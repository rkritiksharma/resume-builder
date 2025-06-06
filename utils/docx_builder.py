from docx import Document
import io

TEMPLATE_PATH = "Resume_Template.docx"

def build_resume_docx(resume_data, selected_projects):
    doc = Document(TEMPLATE_PATH)
    for para in doc.paragraphs:
        if "{{TITLE}}" in para.text:
            para.text = para.text.replace("{{TITLE}}", resume_data["title"])
        if "{{SUMMARY}}" in para.text:
            para.text = para.text.replace("{{SUMMARY}}", resume_data["summary"])
        if "{{SKILLS}}" in para.text:
            para.text = para.text.replace("{{SKILLS}}", ", ".join(resume_data["skills"]))

    for i, para in enumerate(doc.paragraphs):
        if "{{PROJECTS}}" in para.text:
            doc.paragraphs[i].clear()
            for p in selected_projects:
                para = doc.paragraphs[i].insert_paragraph_before()
                title_line = f"{p['title']} ({p.get('client', '-')}, {p.get('duration', '-')})"
                para.add_run(f"• {title_line}\n").bold = True
                if p.get("responsibilities"):
                    for r in p["responsibilities"]:
                        para = doc.paragraphs[i].insert_paragraph_before()
                        para.add_run(f"    - {r}")
                if p.get("url"):
                    para = doc.paragraphs[i].insert_paragraph_before()
                    para.add_run(f"    URL: {p['url']}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def build_cover_letter_docx(cover_letter_text):
    doc = Document()
    for para in cover_letter_text.split("\n\n"):
        doc.add_paragraph(para.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()